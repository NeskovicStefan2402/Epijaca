from flaskApp.models.StavkaRacuna import StavkaRacunaModel
from flaskApp.models.Racun import RacunModel
from flaskApp.other.OstaleFunkcije import OstaleFunkcije
from flaskApp.models.Korisnik import KorisnikModel
from flaskApp.models.Proizvod import ProizvodModel
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
from datetime import date
from flask import jsonify

class Racun:

    @classmethod
    def dodajRacun(cls,data):
        racun=RacunModel(data['idKorisnika'],data['idGosta'])
        try:
            racun.add()
        except:
            return {'Greska':'Problem prilikom cuvanja racuna u bazi!'}

        for i in data['stavke']:
            stavka=StavkaRacunaModel(i['kolicina'],i['idProizvoda'],racun.id)
            try:
                stavka.add()
                
            except:
                return {'Greska':'Problem prilikom cuvanja stavke racuna u bazi!'}
        zaKorisnika=Racun.kreirajSadrzajKorisniku(data,racun.id)
        zaGosta=Racun.kreirajSadrzajGostu(data,racun.id)
        OstaleFunkcije.posaljiEmail(zaKorisnika)
        OstaleFunkcije.posaljiEmail(zaGosta)
        return {'Poruka':'Racun je uspesno zaveden i porudzbina je poslata!'}
    
    @classmethod
    def kreirajSadrzajGostu(cls,data,sifra):   
        korisnik=KorisnikModel.find_one_id(data['idKorisnika'])
        gost=KorisnikModel.find_one_id(data['idGosta'])
        return {'to':gost.email,
                        'subject':'Potvrda porudzbine!',
                        'file':'flaskApp/prijemnice/prijemnice{}.docx'.format(sifra),
                        'poruka':'''Postovani korisnice {0} {1},\n
                                    Vasa porudzbina pod sifrom #{2} je dostavljena korisniku {3} {4} !\n
                                    Pogledajte celokupan asortiman zdravih proizvoda na Epijaci!\n
                                    Pozdrav od Epijaca tima '''.format(gost.ime,gost.prezime,sifra,korisnik.ime,korisnik.prezime)}
    
    @classmethod
    def kreirajSadrzajKorisniku(cls,data,sifra):   
        korisnik=KorisnikModel.find_one_id(data['idKorisnika'])
        gost=KorisnikModel.find_one_id(data['idGosta'])
        stavke=''
        for i in data['stavke']:
            stavke+='\t\t'+ProizvodModel.find_one_idProizvoda(i['idProizvoda']).naziv+','+str(i['kolicina'])+' kom. \n'
        
        poruka='''Postovani korisnice {0} {1},\n
                Vasoj firmi {2} dostavljena je porudzbina pod sifrom #{3} od korisnika {4} {5} !\n
                Molimo Vas da u sto kracem roku dostavite narucene proizvode i prosledite ih na adresu {6}\n.
                Porucene stavke su prosledjene u prilogu. \n
                Pozdrav od Epijaca tima '''.format(korisnik.ime,korisnik.prezime,korisnik.firma,sifra,gost.ime,gost.prezime,gost.adresa,stavke)
        Racun.kreirajPrijemnicu(korisnik,gost,sifra,data['stavke'])
        return {'to':korisnik.email,
                'file':'flaskApp/prijemnice/prijemnice{}.docx'.format(sifra),
                'subject': 'Prijem porudzbine',
                'poruka':poruka}
                            
    @classmethod    
    def kreirajPrijemnicu(cls,korisnik,gost,sifra,stavke):
        document = Document()
        document.add_heading('Prijemnica #{}'.format(sifra), 0)
        document.add_paragraph('Prodavac: {0} {1}'.format(korisnik.ime,korisnik.prezime))
        document.add_paragraph('Kupac: {0} {1}'.format(gost.ime,gost.prezime))
        document.add_paragraph('Vreme kreiranja: {0}'.format(date.today()))
            # p.add_run('bold').bold = True
            # p.add_run(' and some ')
            # p.add_run('italic.').italic = True
        document.add_heading('Stavke racuna', level=1)
            # document.add_paragraph(
            #     'first item in unordered list', style='List Bullet'
            # )
            # document.add_paragraph(
            #     'first item in ordered list', style='List Number'
            # )
        # document.add_picture('flaskApp/images/logo.jpg', width=Inches(1.25))
        records = []
        br=1 
        suma=0
        for i in stavke:
            proizvod=ProizvodModel.find_one_idProizvoda(i['idProizvoda'])
            records.append([br,proizvod.naziv,i['kolicina']])
            suma+=proizvod.cena * i['kolicina']
            br+=1
        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Id'
        hdr_cells[1].text = 'Naziv'
        hdr_cells[2].text = 'Kolicina'
        for id, naziv, kolicina in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(id)
            row_cells[1].text = naziv
            row_cells[2].text = str(kolicina)
        document.add_paragraph('Ukupna suma racuna: {0}'.format(suma))
        document.add_paragraph('POzdrav tvoj Epijaca tim!')
        document.add_page_break()
        naziv='flaskApp/prijemnice/prijemnice'+str(sifra)+'.docx'
        document.save(naziv)


